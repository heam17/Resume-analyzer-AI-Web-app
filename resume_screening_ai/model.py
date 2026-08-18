import os
import re
from collections import Counter

from docx import Document
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity


# -------------------------------------------------------------------
# Skill normalization
# -------------------------------------------------------------------

def normalize_text(text):
    """Normalize text for matching while preserving useful words."""
    text = text.lower()
    text = text.replace("&", " and ")
    text = re.sub(r"[^a-z0-9+#.\-/ ]+", " ", text)
    text = re.sub(r"\s+", " ", text)
    return text.strip()


def skill_variants(skill):
    """
    Generate simple matching variants for common technical skills.
    This is intentionally rule-based and transparent for an
    intermediate-level ML project.
    """
    s = normalize_text(skill)

    variants = {s}

    aliases = {
        "machine learning": ["machine learning", "ml"],
        "artificial intelligence": ["artificial intelligence", "ai"],
        "javascript": ["javascript", "java script", "js"],
        "typescript": ["typescript", "ts"],
        "c++": ["c++", "cpp"],
        "c#": ["c#", "csharp"],
        "node.js": ["node.js", "nodejs", "node js"],
        "react.js": ["react.js", "reactjs", "react js"],
        "react": ["react", "react.js", "reactjs"],
        "angular": ["angular", "angularjs"],
        "spring boot": ["spring boot", "springboot"],
        "spring": ["spring", "spring framework"],
        "sql": ["sql"],
        "postgresql": ["postgresql", "postgres"],
        "mongodb": ["mongodb", "mongo db"],
        "aws": ["aws", "amazon web services"],
        "gcp": ["gcp", "google cloud"],
        "azure": ["azure", "microsoft azure"],
        "docker": ["docker"],
        "kubernetes": ["kubernetes", "k8s"],
        "git": ["git", "github", "gitlab"],
    }

    if s in aliases:
        variants.update(aliases[s])

    return list(variants)


def contains_skill(text, skill):
    """Check whether a skill or one of its common variants occurs."""
    normalized = normalize_text(text)

    for variant in skill_variants(skill):
        if not variant:
            continue

        # Use word boundaries for ordinary words; literal search for
        # symbols such as C++ where boundaries are less reliable.
        if re.search(r"(?<![a-z0-9])" + re.escape(variant) + r"(?![a-z0-9])", normalized):
            return True

    return False


# -------------------------------------------------------------------
# DOCX extraction
# -------------------------------------------------------------------

def extract_text(file_path):
    """
    Extract paragraph and table text from a DOCX resume.
    Tables are included because resume templates often store
    skills/education in tables.
    """
    document = Document(file_path)
    parts = []

    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if text:
            parts.append(text)

    for table in document.tables:
        for row in table.rows:
            cells = []
            for cell in row.cells:
                cell_text = cell.text.strip()
                if cell_text:
                    cells.append(cell_text)
            if cells:
                parts.append(" ".join(cells))

    return "\n".join(parts)


def load_resumes(folder):
    """Load all DOCX resumes from the selected domain folder."""
    resumes = []
    filenames = []

    if not os.path.isdir(folder):
        return resumes, filenames

    for filename in sorted(os.listdir(folder)):
        if not filename.lower().endswith(".docx"):
            continue

        path = os.path.join(folder, filename)

        try:
            text = extract_text(path)

            if text.strip():
                resumes.append(text)
                filenames.append(filename)

        except Exception as exc:
            print(f"Could not process {filename}: {exc}")

    return resumes, filenames


# -------------------------------------------------------------------
# Section detection
# -------------------------------------------------------------------

SECTION_ALIASES = {
    "experience": [
        "professional experience",
        "work experience",
        "experience",
        "employment history",
        "career history",
        "professional background",
    ],
    "education": [
        "education",
        "academic background",
        "academic qualification",
        "qualifications",
    ],
    "certifications": [
        "certifications",
        "certification",
        "licenses and certifications",
        "professional certifications",
    ],
    "projects": [
        "projects",
        "project experience",
        "key projects",
        "project",
    ],
}


def find_section(text, section_name):
    """
    Extract an approximate section from a resume.
    Resume formats vary, so this is deliberately tolerant.
    """
    lines = [line.strip() for line in text.splitlines() if line.strip()]
    aliases = SECTION_ALIASES[section_name]

    start = None

    for i, line in enumerate(lines):
        cleaned = normalize_text(line).strip(" :-")
        if cleaned in aliases:
            start = i + 1
            break

    if start is None:
        return ""

    next_sections = []
    all_aliases = set(
        alias
        for values in SECTION_ALIASES.values()
        for alias in values
    )

    for i in range(start, len(lines)):
        cleaned = normalize_text(lines[i]).strip(" :-")
        if cleaned in all_aliases:
            next_sections.append(i)

    end = min(next_sections) if next_sections else len(lines)

    return "\n".join(lines[start:end])


# -------------------------------------------------------------------
# Experience estimation
# -------------------------------------------------------------------

def estimate_experience_years(text):
    """
    Estimate explicit years of experience from phrases such as:
    '8+ years', '5 years of experience', etc.

    This is not intended to reconstruct exact employment history.
    It is a transparent feature for the internship-level project.
    """
    normalized = normalize_text(text)

    patterns = [
        r"(\d+(?:\.\d+)?)\s*\+\s*years?\s+(?:of\s+)?(?:professional\s+)?experience",
        r"(\d+(?:\.\d+)?)\s+years?\s+(?:of\s+)?(?:professional\s+)?experience",
        r"(\d+(?:\.\d+)?)\s+years?\s+experience",
    ]

    values = []

    for pattern in patterns:
        for match in re.findall(pattern, normalized):
            try:
                values.append(float(match))
            except ValueError:
                pass

    return max(values) if values else 0.0


# -------------------------------------------------------------------
# Feature scoring
# -------------------------------------------------------------------

def exact_skill_score(text, required_skills):
    """
    Percentage of HR-required skills explicitly found in the resume.
    """
    if not required_skills:
        return 0.0, [], []

    matched = []
    missing = []

    for skill in required_skills:
        if contains_skill(text, skill):
            matched.append(skill)
        else:
            missing.append(skill)

    score = (len(matched) / len(required_skills)) * 100
    return score, matched, missing


def tfidf_relevance_score(resume_text, required_skills):
    """
    Semantic-ish lexical relevance using TF-IDF + cosine similarity.
    This complements exact skill matching rather than replacing it.
    """
    query = " ".join(required_skills)

    if not resume_text.strip() or not query.strip():
        return 0.0

    try:
        vectorizer = TfidfVectorizer(
            stop_words="english",
            ngram_range=(1, 2)
        )

        matrix = vectorizer.fit_transform(
            [resume_text, query]
        )

        similarity = cosine_similarity(
            matrix[0:1],
            matrix[1:2]
        )[0][0]

        return float(similarity * 100)

    except ValueError:
        return 0.0


def section_skill_score(section_text, required_skills):
    """
    Skill coverage inside a particular section.
    Used to reward skills appearing in actual work/project content.
    """
    if not section_text or not required_skills:
        return 0.0

    matched = sum(
        1 for skill in required_skills
        if contains_skill(section_text, skill)
    )

    return (matched / len(required_skills)) * 100


def certification_score(text, required_skills):
    """
    Certification relevance score.

    Certifications are not required to exist in every resume.
    We only give certification credit when a required skill is
    explicitly found in the certification section.
    """
    certification_section = find_section(text, "certifications")

    if not certification_section:
        return 0.0

    matched = sum(
        1 for skill in required_skills
        if contains_skill(certification_section, skill)
    )

    if not required_skills:
        return 0.0

    return (matched / len(required_skills)) * 100


def education_score(text):
    """
    Basic education-presence feature.

    We intentionally do not assume a specific degree because the
    current HR input only contains skills.
    """
    education = find_section(text, "education")

    if education:
        return 100.0

    # Some resumes may not have a clean heading but still contain
    # common degree terms.
    normalized = normalize_text(text)

    degree_terms = [
        "bachelor",
        "master",
        "b.tech",
        "b.e.",
        "bca",
        "mca",
        "m.tech",
        "mba",
        "ph.d",
        "university",
        "college",
    ]

    return 50.0 if any(term in normalized for term in degree_terms) else 0.0


def experience_score(text, required_skills):
    """
    Experience score combines:
    - presence of an experience section
    - explicit years of experience
    - required-skill coverage in experience text

    Since HR has not yet entered a minimum years requirement,
    years are used as a supporting signal, not a hard rejection.
    """
    experience_section = find_section(text, "experience")
    years = estimate_experience_years(text)

    section_present = 100.0 if experience_section else 0.0
    relevant_experience = section_skill_score(
        experience_section,
        required_skills
    )

    # Cap the years component at 10 years so long resumes do not
    # dominate the score simply because they are older.
    years_component = min(years / 10.0, 1.0) * 100.0

    if experience_section:
        score = (
            relevant_experience * 0.60
            + years_component * 0.25
            + section_present * 0.15
        )
    else:
        score = years_component * 0.50 + section_present * 0.50

    return min(score, 100.0), years


# -------------------------------------------------------------------
# Final candidate evaluation
# -------------------------------------------------------------------

def evaluate_resume(text, filename, required_skills):
    """
    Calculate a transparent weighted candidate match score.

    Current weights:
      Skills                  40%
      Relevant Experience     25%
      Education               10%
      Relevant Work/Projects  15%
      Certifications          10%

    The score is a MATCH SCORE, not statistical model accuracy.
    """
    skills_score, matched_skills, missing_skills = exact_skill_score(
        text,
        required_skills
    )

    # TF-IDF provides contextual relevance. Blend it lightly into
    # the skills feature rather than allowing generic words to dominate.
    lexical_relevance = tfidf_relevance_score(
        text,
        required_skills
    )

    final_skills_score = (
        skills_score * 0.75
        + lexical_relevance * 0.25
    )

    experience_section = find_section(text, "experience")
    projects_section = find_section(text, "projects")

    experience, years = experience_score(
        text,
        required_skills
    )

    # Work + projects: prefer actual experience content where present.
    work_relevance = section_skill_score(
        experience_section,
        required_skills
    )

    project_relevance = section_skill_score(
        projects_section,
        required_skills
    )

    if experience_section and projects_section:
        work_project_score = work_relevance * 0.70 + project_relevance * 0.30
    elif experience_section:
        work_project_score = work_relevance
    elif projects_section:
        work_project_score = project_relevance
    else:
        work_project_score = 0.0

    education = education_score(text)
    certifications = certification_score(text, required_skills)

    # Weighted final score.
    final_score = (
        final_skills_score * 0.40
        + experience * 0.25
        + education * 0.10
        + work_project_score * 0.15
        + certifications * 0.10
    )

    return {
        "filename": filename,
        "score": round(min(final_score, 100.0), 2),
        "skills_score": round(final_skills_score, 2),
        "experience_score": round(experience, 2),
        "education_score": round(education, 2),
        "work_project_score": round(work_project_score, 2),
        "certification_score": round(certifications, 2),
        "experience_years": years,
        "matched_skills": matched_skills,
        "missing_skills": missing_skills,
    }


def evaluate_resumes(resumes, filenames, required_skills):
    """Evaluate and rank all resumes."""
    results = []

    for text, filename in zip(resumes, filenames):
        result = evaluate_resume(
            text,
            filename,
            required_skills
        )
        results.append(result)

    results.sort(
        key=lambda item: item["score"],
        reverse=True
    )

    for rank, result in enumerate(results, start=1):
        result["rank"] = rank

    return results
